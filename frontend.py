import streamlit as st
import requests
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from io import BytesIO
from docx import Document

# SESSION STATE INIT
if "usage_count" not in st.session_state:
    st.session_state.usage_count = 0

if "analysis_result" not in st.session_state:
    st.session_state.analysis_result = None

if "optimized_resume" not in st.session_state:
    st.session_state.optimized_resume = None

if "cover_letter" not in st.session_state:
    st.session_state.cover_letter = None


# PDF Generator
def generate_pdf_report(title, content):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph(title, styles["Title"]))
    elements.append(Spacer(1, 0.3 * inch))
    elements.append(Paragraph(content.replace("\n", "<br/>"), styles["Normal"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer

# DOCX Generator
def generate_docx_resume(content):
    document = Document()
    document.add_heading("Document", level=1)

    for line in content.split("\n"):
        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# PAGE CONFIG
st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="🧠",
    layout="centered"
)

# SIDEBAR
with st.sidebar:
    st.title("⚙️ Settings")
    theme_mode = st.toggle("🌙 Dark Mode")

if theme_mode:
    st.markdown("""
        <style>
        .stApp { background-color: #0E1117; color: white; }
        textarea, input { background-color: #262730 !important; color: white !important; }
        </style>
    """, unsafe_allow_html=True)

# AUTH
PASSWORD = "resume123"
password_input = st.text_input("Enter Access Password", type="password")

if password_input != PASSWORD:
    st.warning("Please enter valid password to use the platform.")
    st.stop()

# MAIN UI
st.title("🧠 AI Resume Analyzer")
st.write("Upload your resume (PDF) and get instant AI feedback.")

uploaded_file = st.file_uploader("Upload Resume", type=["pdf"])
jd_input = st.text_area("Paste Job Description (Optional)")


# ANALYZE BUTTON
if uploaded_file and st.button("🔍 Analyze Resume"):

    with st.spinner("Analyzing resume..."):

        form_data = {}
        if jd_input.strip():
            form_data["job_description"] = jd_input

        response = requests.post(
            "http://backend:8000/analyze-resume",
            files={
                "file": (
                    uploaded_file.name,
                    uploaded_file.getvalue(),
                    "application/pdf"
                )
            },
            data=form_data
        )

        if response.status_code == 200:
            st.session_state.analysis_result = response.json()
            st.session_state.usage_count += 1
        else:
            st.error("Backend error occurred.")


# SHOW ANALYSIS (PERSISTENT)
if st.session_state.analysis_result:

    result = st.session_state.analysis_result

    st.success("✅ Analysis Complete!")
    st.caption(f"📈 Analyses this session: {st.session_state.usage_count}")

   
    # ATS SECTION
    if "score" in result:
        st.subheader("📊 ATS Score")
        st.progress(result["score"] / 100)
        st.write(f"Score: {result['score']}/100")

        st.subheader("Strengths")
        st.write(result.get("strengths", "N/A"))

        st.subheader("Improvements")
        st.write(result.get("improvements", "N/A"))

        st.subheader("Overall Summary")
        st.write(result.get("overall_summary", "N/A"))

        st.markdown("---")

    # JD MATCH SECTION
    if "match_score" in result:
        st.subheader("🎯 Resume vs JD Match Score")
        st.progress(result["match_score"] / 100)
        st.write(f"Match Score: {result['match_score']}/100")

        matched = result.get("matched_skills", [])
        missing = result.get("missing_skills", [])

        if isinstance(matched, str):
            matched = [s.strip() for s in matched.split(",") if s.strip()]

        if isinstance(missing, str):
            missing = [s.strip() for s in missing.split(",") if s.strip()]

        col1, col2 = st.columns(2)
        col1.metric("Matched Skills", len(matched))
        col2.metric("Missing Skills", len(missing))

        if len(matched) + len(missing) > 0:
            percent = int((len(matched) / (len(matched) + len(missing))) * 100)
            st.progress(percent / 100)
            st.write(f"Skill Coverage: **{percent}%**")

        if missing:
            st.markdown("### 🚨 High Priority Skills Missing")
            for skill in missing[:5]:
                st.warning(skill)

        st.subheader("Improvement Suggestions")
        st.write(result.get("improvement_suggestions", "N/A"))

        st.markdown("---")

   
    # OPTIMIZE BUTTON
    if st.button("🚀 Generate Optimized Resume"):

        with st.spinner("Optimizing resume..."):

            optimize_response = requests.post(
                "http://backend:8000/optimize-resume",
                files={
                    "file": (
                        uploaded_file.name,
                        uploaded_file.getvalue(),
                        "application/pdf"
                    )
                }
            )

            if optimize_response.status_code == 200:
                st.session_state.optimized_resume = optimize_response.json()["optimized_resume"]
            else:
                st.error("Optimization failed.")

    # SHOW OPTIMIZED
    if st.session_state.optimized_resume:
        st.subheader("✨ Optimized Resume")
        st.write(st.session_state.optimized_resume)

        docx_buffer = generate_docx_resume(st.session_state.optimized_resume)

        st.download_button(
            "📄 Download Optimized Resume (.docx)",
            data=docx_buffer,
            file_name="optimized_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.markdown("---")

    
    # COVER LETTER
    if jd_input.strip() and st.button("✨ Generate Cover Letter"):

        with st.spinner("Generating cover letter..."):

            cover_response = requests.post(
                "http://backend:8000/generate-cover-letter",
                files={
                    "file": (
                        uploaded_file.name,
                        uploaded_file.getvalue(),
                        "application/pdf"
                    )
                },
                data={"job_description": jd_input}
            )

            if cover_response.status_code == 200:
                st.session_state.cover_letter = cover_response.json()["cover_letter"]
            else:
                st.error("Cover letter generation failed.")

    # SHOW COVER LETTER
    if st.session_state.cover_letter:
        st.subheader("📝 Generated Cover Letter")
        st.write(st.session_state.cover_letter)

        docx_buffer = generate_docx_resume(st.session_state.cover_letter)

        st.download_button(
            "📄 Download Cover Letter (.docx)",
            data=docx_buffer,
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
